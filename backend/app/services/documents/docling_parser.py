from __future__ import annotations

from pathlib import Path

from app.core.constants import SUPPORTED_EXTENSIONS
from app.core.exceptions import DocumentParsingError
from app.schemas.document import DocumentKind, DocumentText
from app.services.documents.section_segmenter import segment_sections


class DoclingParser:
    def extract(self, path: str | Path, kind: DocumentKind = DocumentKind.unknown) -> DocumentText:
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise DocumentParsingError(f"Format non supporte: {suffix}")
        if suffix == ".pdf":
            text = self._extract_pdf(path)
        elif suffix == ".docx":
            text = self._extract_docx(path)
        else:
            text = self._extract_text(path)
        text = text.strip()
        if not text:
            raise DocumentParsingError(f"Aucun texte exploitable extrait depuis {path.name}.")
        return DocumentText(filename=path.name, kind=kind, text=text, char_count=len(text), sections=segment_sections(text))

    def _extract_pdf(self, path: Path) -> str:
        import fitz
        parts = []
        with fitz.open(path) as doc:
            for page in doc:
                page_text = page.get_text("text")
                if page_text.strip():
                    parts.append(page_text)
        return "\n".join(parts)

    def _extract_docx(self, path: Path) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _extract_text(self, path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")

